"""
File parser service — extract text from uploaded resume files (PDF / DOCX)
and pass the result through the AI parsing pipeline.

Dependencies (already in requirements.txt):
  PyMuPDF  (import fitz)  — PDF parsing
  python-docx             — DOCX parsing
"""

from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)

# Supported MIME types
_MIME_PDF = "application/pdf"
_MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_MIME_DOC = "application/msword"

# Minimum extracted text length to consider parsing successful
_MIN_TEXT_LENGTH = 50


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

async def parse_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF file.

    Uses PyMuPDF (fitz) for reliable text extraction including multi-column
    layouts. Pages are joined with double newlines.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        Extracted text as a single string, or empty string on failure.
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages: list[str] = []
        for page in doc:
            text = page.get_text("text")  # type: ignore[attr-defined]
            if text.strip():
                pages.append(text.strip())
        doc.close()
        return "\n\n".join(pages)
    except ImportError:
        logger.error("PyMuPDF (fitz) is not installed. Install with: pip install PyMuPDF")
        raise
    except Exception as exc:
        logger.error("Failed to parse PDF: %s", exc)
        return ""


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

async def parse_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file.

    Iterates over all paragraphs (including those inside tables) and joins
    non-empty lines. Table cells are separated by tabs for readability.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Extracted text as a single string, or empty string on failure.
    """
    try:
        from docx import Document  # python-docx

        doc = Document(io.BytesIO(file_bytes))
        lines: list[str] = []

        # Main paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Table cells (resumes sometimes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cell_texts:
                    lines.append("\t".join(cell_texts))

        return "\n".join(lines)
    except ImportError:
        logger.error("python-docx is not installed. Install with: pip install python-docx")
        raise
    except Exception as exc:
        logger.error("Failed to parse DOCX: %s", exc)
        return ""


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def _detect_mime_from_bytes(file_bytes: bytes) -> str | None:
    """Heuristic MIME detection by magic bytes when MIME type is unknown."""
    if file_bytes[:4] == b"%PDF":
        return _MIME_PDF
    # DOCX / ZIP (DOCX is a ZIP archive)
    if file_bytes[:2] == b"PK":
        return _MIME_DOCX
    return None


async def extract_resume_data(file_bytes: bytes, mime_type: str) -> dict:
    """Parse an uploaded resume file and return structured profile data.

    Workflow:
      1. Detect file type from mime_type (with magic-byte fallback).
      2. Extract raw text (parse_pdf or parse_docx).
      3. Validate that enough text was extracted.
      4. Call ai_service.parse_resume_file(text) to get structured data.

    Args:
        file_bytes: Raw file bytes as received from Telegram.
        mime_type:  MIME type string reported by Telegram.

    Returns:
        Structured dict with keys: name, contacts, summary,
        work_experiences, skills, education, languages, certifications.

    Raises:
        ValueError: If the file type is unsupported or text extraction failed.
    """
    # Lazy import to avoid circular dependency at module load time
    from bot.services.ai_service import parse_resume_file

    # --- Normalise MIME type ---
    normalised_mime = (mime_type or "").lower().strip()

    # Fallback to magic-byte detection
    if normalised_mime not in (_MIME_PDF, _MIME_DOCX, _MIME_DOC):
        detected = _detect_mime_from_bytes(file_bytes)
        if detected:
            logger.info(
                "MIME type %r not recognised; detected %r from magic bytes.",
                mime_type,
                detected,
            )
            normalised_mime = detected
        else:
            raise ValueError(
                f"Unsupported file type: {mime_type!r}. "
                "Please upload a PDF or DOCX file."
            )

    # --- Extract text ---
    if normalised_mime == _MIME_PDF:
        text = await parse_pdf(file_bytes)
    else:  # DOCX or DOC
        text = await parse_docx(file_bytes)

    if len(text.strip()) < _MIN_TEXT_LENGTH:
        raise ValueError(
            "Не удалось извлечь текст из файла. "
            "Убедитесь, что файл не защищён паролем и содержит текст (не сканированное изображение)."
        )

    logger.info("Extracted %d characters from uploaded resume file.", len(text))

    # --- AI parsing ---
    try:
        structured = await parse_resume_file(text)
    except Exception as exc:
        logger.error("AI resume parsing failed: %s", exc)
        # Return a minimal dict so the caller can handle gracefully
        return {
            "name": "",
            "contacts": {},
            "summary": "",
            "work_experiences": [],
            "skills": [],
            "education": [],
            "languages": [],
            "certifications": [],
            "_raw_text": text,
            "_parse_error": str(exc),
        }

    # Attach raw text for debugging / fallback use
    structured["_raw_text"] = text
    return structured
